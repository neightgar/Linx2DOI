"""
Главное окно приложения
"""

import sys
import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QGridLayout,
    QPushButton, QLabel, QProgressBar, QTextEdit, QFileDialog,
    QMessageBox, QFrame, QLineEdit, QTabWidget, QGroupBox,
    QTableWidget, QTableWidgetItem, QHeaderView, QCheckBox, QAbstractItemView
)
from PyQt6.QtCore import Qt, pyqtSlot, QSettings
from PyQt6.QtGui import QColor, QFont, QBrush, QIcon

from ..core.config import SETTINGS_ORG, SETTINGS_APP
from ..core.worker import WorkerThread


def get_icon_path():
    """Получает путь к иконке для dev и compiled режимов"""
    # Возможные расположения иконки
    if getattr(sys, 'frozen', False):
        # Режим compiled (Nuitka/PyInstaller)
        base_path = Path(sys.executable).parent
    else:
        # Режим разработки
        base_path = Path(__file__).parent.parent

    # Проверяем возможные расположения
    icon_locations = [
        base_path / "icon.ico",
        base_path / "src" / "icon.ico",
        Path(__file__).parent.parent / "icon.ico",
    ]

    for icon_path in icon_locations:
        if icon_path.exists():
            return icon_path

    return None
from ..core.ris_exporter import RISExporter
from .delegates import HyperlinkDelegate, ManualDOIDelegate


# Темная палитра для приложения
class AppColors:
    # Фоновые цвета (темная тема)
    BG_DARKEST = "#1e1e20"         # Самый темный фон
    BG_DARK = "#2d2d30"            # Темный фон
    BG_MEDIUM = "#3e3e42"          # Средний фон
    BG_LIGHT = "#4d4d52"           # Светлый фон

    # Текст
    TEXT_PRIMARY = "#e6e6e6"       # Основной текст
    TEXT_SECONDARY = "#969696"     # Вторичный текст
    TEXT_DISABLED = "#7f7f7f"      # Отключенный текст

    # Акценты (синие оттенки)
    PRIMARY = "#2a82da"            # Основной синий
    PRIMARY_DARK = "#1e5f99"       # Темный синий
    PRIMARY_LIGHT = "#3d95ed"      # Светлый синий

    # Статусы
    SUCCESS = "#4ec9b0"            # Зеленый (cyan)
    WARNING = "#ce9178"            # Оранжевый
    ERROR = "#f48771"              # Красный
    INFO = "#2a82da"               # Голубой

    # Границы и разделители
    BORDER = "#555558"             # Границы
    BORDER_LIGHT = "#656569"       # Светлые границы

    # Таблица
    TABLE_HEADER = "#3e3e42"       # Заголовок таблицы
    TABLE_ROW = "#2d2d30"          # Строка таблицы
    TABLE_ROW_ALT = "#333337"      # Альтернативная строка
    TABLE_SELECTED = "#264f78"     # Выбранная строка

    # Специальные
    SCROLLBAR = "#424245"          # Полоса прокрутки
    HOVER = "#3e3e42"              # Hover эффект


# Светлая палитра для приложения
class LightColors:
    # Фоновые цвета (светлая тема)
    BG_DARKEST = "#e8e8e8"         # Для заголовка
    BG_LIGHTEST = "#ffffff"        # Белый фон
    BG_LIGHT = "#f5f5f5"           # Светло-серый
    BG_MEDIUM = "#e0e0e0"          # Средне-серый
    BG_DARK = "#d0d0d0"            # Темнее

    # Текст
    TEXT_PRIMARY = "#2d2d30"       # Темный текст
    TEXT_SECONDARY = "#616161"     # Серый текст
    TEXT_DISABLED = "#9e9e9e"      # Отключенный текст

    # Акценты (синие оттенки)
    PRIMARY = "#2a82da"            # Основной синий
    PRIMARY_DARK = "#1e5f99"       # Темный синий
    PRIMARY_LIGHT = "#3d95ed"      # Светлый синий

    # Статусы
    SUCCESS = "#4caf50"            # Зеленый
    WARNING = "#ff9800"            # Оранжевый
    ERROR = "#f44336"              # Красный
    INFO = "#2196f3"               # Голубой

    # Границы и разделители
    BORDER = "#cccccc"             # Границы
    BORDER_LIGHT = "#e0e0e0"       # Светлые границы

    # Таблица
    TABLE_HEADER = "#e0e0e0"       # Заголовок таблицы
    TABLE_ROW = "#ffffff"          # Строка таблицы
    TABLE_ROW_ALT = "#f9f9f9"      # Альтернативная строка
    TABLE_SELECTED = "#bbdefb"     # Выбранная строка

    # Специальные
    SCROLLBAR = "#cccccc"          # Полоса прокрутки
    HOVER = "#eeeeee"              # Hover эффект


class MainWindow(QMainWindow):
    """Главное окно приложения"""

    def __init__(self):
        super().__init__()
        self.setWindowTitle("Linx2DOI - DOI Extractor & APA Formatter")
        self.resize(1300, 850)
        self.setMinimumSize(1100, 700)

        # Установка иконки приложения
        icon_path = get_icon_path()
        if icon_path:
            self.setWindowIcon(QIcon(str(icon_path)))

        # Состояние приложения
        self.input_path = None
        self.output_path = None
        self.settings = QSettings(SETTINGS_ORG, SETTINGS_APP)
        self.all_items = []
        self.current_items = []
        self.table_items_map = {}
        self.worker = None

        # Загрузка темы из настроек (по умолчанию светлая)
        self.current_theme = self.settings.value("theme", "light")
        self.colors = LightColors if self.current_theme == "light" else AppColors

        # Применяем стили текущей темы
        self.apply_theme()
        self.init_ui()

    def apply_theme(self):
        """Применение стилей текущей темы"""
        self.setStyleSheet(f"""
            QMainWindow {{
                background-color: {self.colors.BG_LIGHT};
            }}

            QLabel {{
                color: {self.colors.TEXT_PRIMARY};
            }}

            QPushButton {{
                background-color: {self.colors.BG_MEDIUM};
                color: {self.colors.TEXT_PRIMARY};
                border: 1px solid {self.colors.BORDER};
                border-radius: 4px;
                padding: 8px 16px;
                font-weight: 600;
                font-size: 12px;
            }}
            QPushButton:hover {{
                background-color: {self.colors.BG_LIGHT};
                border-color: {self.colors.BORDER_LIGHT};
            }}
            QPushButton:pressed {{
                background-color: {self.colors.BG_DARKEST};
            }}
            QPushButton:disabled {{
                background-color: {self.colors.BG_DARKEST};
                color: {self.colors.TEXT_DISABLED};
                border-color: {self.colors.BORDER};
            }}

            QLineEdit {{
                background-color: {self.colors.BG_DARKEST};
                color: {self.colors.TEXT_PRIMARY};
                border: 1px solid {self.colors.BORDER};
                border-radius: 4px;
                padding: 6px 10px;
                font-size: 13px;
                selection-background-color: {self.colors.PRIMARY};
            }}
            QLineEdit:focus {{
                border-color: {self.colors.PRIMARY};
            }}

            QProgressBar {{
                border: none;
                border-radius: 3px;
                background-color: {self.colors.BG_DARKEST};
                height: 6px;
                text-align: center;
                color: {'#2d2d30' if self.current_theme == 'light' else '#ffffff'};
                font-weight: 600;
            }}
            QProgressBar::chunk {{
                background-color: {self.colors.PRIMARY};
                border-radius: 3px;
            }}

            QTabWidget::pane {{
                border: 1px solid {self.colors.BORDER};
                border-radius: 4px;
                background-color: {self.colors.BG_DARK};
                padding: 10px;
            }}

            QTabBar::tab {{
                background-color: {self.colors.BG_MEDIUM};
                color: {self.colors.TEXT_SECONDARY};
                padding: 8px 20px;
                border: 1px solid {self.colors.BORDER};
                border-bottom: none;
                border-top-left-radius: 4px;
                border-top-right-radius: 4px;
                margin-right: 2px;
                font-weight: 600;
            }}
            QTabBar::tab:selected {{
                background-color: {self.colors.PRIMARY};
                color: white;
            }}
            QTabBar::tab:hover:!selected {{
                background-color: {self.colors.BG_LIGHT};
                color: {self.colors.TEXT_PRIMARY};
            }}

            QGroupBox {{
                background-color: {self.colors.BG_DARK};
                border: 1px solid {self.colors.BORDER};
                border-radius: 6px;
                margin-top: 12px;
                font-weight: 600;
                padding-top: 10px;
                color: {self.colors.TEXT_PRIMARY};
            }}
            QGroupBox::title {{
                color: {self.colors.PRIMARY_LIGHT};
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px;
            }}

            QScrollBar:vertical {{
                background-color: {self.colors.BG_DARK};
                width: 14px;
                border-radius: 7px;
            }}
            QScrollBar::handle:vertical {{
                background-color: {self.colors.SCROLLBAR};
                border-radius: 7px;
                min-height: 30px;
            }}
            QScrollBar::handle:vertical:hover {{
                background-color: {self.colors.BG_LIGHT};
            }}
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {{
                height: 0px;
            }}

            QTextEdit {{
                background-color: {self.colors.BG_DARKEST};
                color: {self.colors.TEXT_PRIMARY};
                border: 1px solid {self.colors.BORDER};
                border-radius: 4px;
                selection-background-color: {self.colors.PRIMARY};
            }}

            QMessageBox {{
                background-color: {self.colors.BG_DARK};
            }}
            QMessageBox QLabel {{
                color: {self.colors.TEXT_PRIMARY};
            }}
        """)

    def init_ui(self):
        """Инициализация пользовательского интерфейса"""
        # Центральный виджет
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Заголовок приложения
        self.create_header(main_layout)

        # Основной контент с отступами
        content_widget = QWidget()
        content_layout = QVBoxLayout(content_widget)
        content_layout.setContentsMargins(15, 15, 15, 15)
        content_layout.setSpacing(12)

        # Панель управления (компактная)
        self.create_control_panel(content_layout)

        # Вкладки: Результаты и Лог
        self.create_tabs(content_layout)

        main_layout.addWidget(content_widget)

    def create_header(self, layout):
        """Создание заголовка приложения"""
        header = QFrame()
        header.setStyleSheet(f"""
            QFrame {{
                background-color: {self.colors.BG_DARKEST};
            }}
        """)
        header.setFixedHeight(70)
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(20, 10, 20, 10)

        # Иконка приложения
        icon_path = get_icon_path()
        if icon_path:
            icon_label = QLabel()
            pixmap = QIcon(str(icon_path)).pixmap(48, 48)
            icon_label.setPixmap(pixmap)
            icon_label.setFixedSize(48, 48)
            header_layout.addWidget(icon_label)

        # Текстовая часть: название и подзаголовок
        text_layout = QVBoxLayout()
        text_layout.setSpacing(2)

        title = QLabel("Linx2DOI")
        title_color = "#ff8c00" if self.current_theme == "light" else "#ffa726"
        title.setStyleSheet(f"""
            QLabel {{
                color: {title_color};
                font-size: 26px;
                font-weight: 700;
                letter-spacing: 1px;
            }}
        """)
        text_layout.addWidget(title)

        subtitle = QLabel("DOI Extractor & APA Bibliography Generator")
        subtitle.setStyleSheet(f"""
            QLabel {{
                color: {self.colors.TEXT_SECONDARY};
                font-size: 12px;
                font-weight: 500;
            }}
        """)
        text_layout.addWidget(subtitle)

        header_layout.addLayout(text_layout)
        header_layout.addStretch()

        # Кнопка переключения темы
        theme_icon = "🌙" if self.current_theme == "light" else "☀️"
        theme_text = "Темная" if self.current_theme == "light" else "Светлая"
        self.theme_btn = QPushButton(f"{theme_icon} {theme_text}")
        self.theme_btn.setStyleSheet("""
            QPushButton {
                padding: 6px 16px;
                font-size: 12px;
            }
        """)
        self.theme_btn.setToolTip("Переключить тему оформления")
        self.theme_btn.clicked.connect(self.toggle_theme)
        header_layout.addWidget(self.theme_btn)

        # Кнопка справки
        help_btn = QPushButton("Справка")
        help_btn.setStyleSheet("""
            QPushButton {
                padding: 6px 16px;
                font-size: 12px;
            }
        """)
        help_btn.clicked.connect(self.show_help)
        header_layout.addWidget(help_btn)

        layout.addWidget(header)

    def create_control_panel(self, layout):
        """Создание компактной панели управления"""
        control_group = QGroupBox("")
        control_layout = QVBoxLayout(control_group)
        control_layout.setSpacing(10)

        # Первая строка: Email и файл
        row1 = QHBoxLayout()
        row1.setSpacing(15)

        # Email
        email_widget = QWidget()
        email_layout = QVBoxLayout(email_widget)
        email_layout.setContentsMargins(0, 0, 0, 0)
        email_layout.setSpacing(4)

        email_label = QLabel("Email для NCBI API:")
        email_label.setStyleSheet(f"color: {self.colors.TEXT_SECONDARY}; font-size: 11px; font-weight: 600;")
        email_layout.addWidget(email_label)

        self.email_input = QLineEdit()
        self.email_input.setPlaceholderText("example@gmail.com")
        self.email_input.setFixedHeight(32)
        saved_email = self.settings.value("user_email", "")
        self.email_input.setText(saved_email)
        email_layout.addWidget(self.email_input)

        row1.addWidget(email_widget, 1)

        # Файл
        file_widget = QWidget()
        file_layout = QVBoxLayout(file_widget)
        file_layout.setContentsMargins(0, 0, 0, 0)
        file_layout.setSpacing(4)

        file_label = QLabel("Документ:")
        file_label.setStyleSheet(f"color: {self.colors.TEXT_SECONDARY}; font-size: 11px; font-weight: 600;")
        file_layout.addWidget(file_label)

        file_row = QHBoxLayout()
        file_row.setSpacing(8)

        self.file_label = QLabel("Файл не выбран")
        self.file_label.setStyleSheet(f"""
            QLabel {{
                background-color: {self.colors.BG_DARKEST};
                border: 1px solid {self.colors.BORDER};
                border-radius: 4px;
                padding: 6px 10px;
                color: {self.colors.TEXT_SECONDARY};
                font-size: 12px;
            }}
        """)
        file_row.addWidget(self.file_label, 1)

        select_file_btn = QPushButton("📁 Выбрать")
        select_file_btn.setFixedHeight(32)
        select_file_btn.clicked.connect(self.select_file)
        file_row.addWidget(select_file_btn)

        file_layout.addLayout(file_row)
        row1.addWidget(file_widget, 2)

        control_layout.addLayout(row1)

        # Вторая строка: Кнопки действий и прогресс
        row2 = QHBoxLayout()
        row2.setSpacing(10)

        # Кнопки действий
        self.process_btn = QPushButton("🚀 Запустить обработку")
        self.process_btn.setFixedHeight(34)
        self.process_btn.clicked.connect(self.start_processing)
        self.process_btn.setEnabled(False)
        row2.addWidget(self.process_btn)

        self.apply_manual_doi_btn = QPushButton("🔍 Обработать выбранные")
        self.apply_manual_doi_btn.setFixedHeight(34)
        self.apply_manual_doi_btn.clicked.connect(self.process_selected)
        self.apply_manual_doi_btn.setEnabled(False)
        self.apply_manual_doi_btn.setToolTip(
            "Обработать записи с галочками ИЛИ ручными DOI\n"
            "• Галочка + DOI вручную = использует ручной DOI\n"
            "• Галочка без DOI = автопоиск\n"
            "• DOI вручную без галочки = использует ручной DOI"
        )
        row2.addWidget(self.apply_manual_doi_btn)

        self.open_btn = QPushButton("📄 Открыть результат")
        self.open_btn.setFixedHeight(34)
        self.open_btn.clicked.connect(self.open_result)
        self.open_btn.setEnabled(False)
        row2.addWidget(self.open_btn)

        self.export_ris_btn = QPushButton("📥 Экспорт в RIS")
        self.export_ris_btn.setFixedHeight(34)
        self.export_ris_btn.clicked.connect(self.export_to_ris)
        self.export_ris_btn.setEnabled(False)
        self.export_ris_btn.setToolTip("Экспорт в формат RIS для Mendeley, Zotero, EndNote")
        row2.addWidget(self.export_ris_btn)

        self.collect_publications_btn = QPushButton("📚 Собрать публикации")
        self.collect_publications_btn.setFixedHeight(34)
        self.collect_publications_btn.clicked.connect(self.collect_publications)
        self.collect_publications_btn.setEnabled(False)
        self.collect_publications_btn.setToolTip(
            "Создать docx файл с публикациями в формате APA\n"
            "Исключаются записи со статусом 'ОТСУТСТВУЕТ В ТЕКСТЕ'"
        )
        row2.addWidget(self.collect_publications_btn)

        row2.addStretch()

        # Прогресс
        progress_widget = QWidget()
        progress_layout = QVBoxLayout(progress_widget)
        progress_layout.setContentsMargins(0, 0, 0, 0)
        progress_layout.setSpacing(2)

        self.progress_label = QLabel("Готов к работе")
        self.progress_label.setStyleSheet(f"""
            QLabel {{
                color: {self.colors.TEXT_SECONDARY};
                font-size: 11px;
                font-weight: 600;
            }}
        """)
        self.progress_label.setAlignment(Qt.AlignmentFlag.AlignRight)
        progress_layout.addWidget(self.progress_label)

        self.progress_bar = QProgressBar()
        self.progress_bar.setFixedWidth(300)
        self.progress_bar.setFixedHeight(22)
        self.progress_bar.setValue(0)
        progress_layout.addWidget(self.progress_bar, alignment=Qt.AlignmentFlag.AlignRight)

        row2.addWidget(progress_widget)

        control_layout.addLayout(row2)

        layout.addWidget(control_group)

    def create_tabs(self, layout):
        """Создание вкладок для результатов и лога"""
        self.tabs = QTabWidget()

        # Вкладка "Результаты"
        results_tab = QWidget()
        results_layout = QVBoxLayout(results_tab)
        results_layout.setContentsMargins(0, 0, 0, 0)
        results_layout.setSpacing(10)

        # Таблица
        self.results_table = QTableWidget()
        self.results_table.setColumnCount(7)
        self.results_table.setHorizontalHeaderLabels(
            ["№", "Название", "Заголовок статьи", "DOI", "Ручной DOI", "Статус", "✓"])

        # Настройка таблицы
        self.results_table.setAlternatingRowColors(True)
        self.results_table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)

        # Настройка режимов растягивания колонок
        header = self.results_table.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.ResizeMode.Fixed)      # №
        header.setSectionResizeMode(1, QHeaderView.ResizeMode.Stretch)    # Название
        header.setSectionResizeMode(2, QHeaderView.ResizeMode.Stretch)    # Заголовок статьи
        header.setSectionResizeMode(3, QHeaderView.ResizeMode.Stretch)    # DOI
        header.setSectionResizeMode(4, QHeaderView.ResizeMode.Stretch)    # Ручной DOI
        header.setSectionResizeMode(5, QHeaderView.ResizeMode.Fixed)      # Статус
        header.setSectionResizeMode(6, QHeaderView.ResizeMode.Fixed)      # Выбрать

        # Минимальные ширины для растягиваемых колонок
        self.results_table.setColumnWidth(0, 50)   # №
        header.setMinimumSectionSize(30)
        self.results_table.setColumnWidth(1, 200)  # Название
        self.results_table.setColumnWidth(2, 200)  # Заголовок статьи
        self.results_table.setColumnWidth(3, 250)  # DOI
        self.results_table.setColumnWidth(4, 250)  # Ручной DOI (увеличено!)
        self.results_table.setColumnWidth(5, 150)  # Статус
        self.results_table.setColumnWidth(6, 50)   # Выбрать

        # Обработчик клика по заголовку колонки "✓" для выбора/снятия всех
        header.sectionClicked.connect(self.on_header_clicked)

        # Делегаты для гиперссылок
        self.hyperlink_delegate = HyperlinkDelegate(self.results_table)
        self.results_table.setItemDelegateForColumn(1, self.hyperlink_delegate)
        self.results_table.setItemDelegateForColumn(2, self.hyperlink_delegate)

        # Делегат для ручного DOI (выравнивание вверх)
        self.manual_doi_delegate = ManualDOIDelegate(self.results_table)
        self.results_table.setItemDelegateForColumn(4, self.manual_doi_delegate)

        self.results_table.setStyleSheet(f"""
            QTableWidget {{
                background-color: {self.colors.TABLE_ROW};
                color: {self.colors.TEXT_PRIMARY};
                border: 1px solid {self.colors.BORDER};
                border-radius: 4px;
                gridline-color: {self.colors.BORDER};
                font-size: 12px;
                selection-background-color: {self.colors.TABLE_SELECTED};
            }}
            QTableWidget::item {{
                padding: 8px 8px;
                border: none;
            }}
            QTableWidget::item:selected {{
                background-color: {self.colors.TABLE_SELECTED};
                color: {'#2d2d30' if self.current_theme == 'light' else 'white'};
                font-weight: {'700' if self.current_theme == 'light' else '400'};
            }}
            QHeaderView::section {{
                background-color: {self.colors.TABLE_HEADER};
                color: {self.colors.TEXT_PRIMARY};
                padding: 10px 4px;
                border: none;
                border-bottom: 2px solid {self.colors.PRIMARY};
                font-weight: 700;
                font-size: 11px;
            }}
            QTableWidget::item:alternate {{
                background-color: {self.colors.TABLE_ROW_ALT};
            }}
        """)

        # Устанавливаем высоту строк для комфортного ввода
        self.results_table.verticalHeader().setDefaultSectionSize(45)

        # Колонка чекбоксов теперь видна для выборочного парсинга
        # self.results_table.setColumnHidden(6, True)

        results_layout.addWidget(self.results_table)

        self.tabs.addTab(results_tab, "Результаты")

        # Вкладка "Лог"
        log_tab = QWidget()
        log_layout = QVBoxLayout(log_tab)
        log_layout.setContentsMargins(0, 0, 0, 0)

        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        self.log_text.setStyleSheet(f"""
            QTextEdit {{
                background-color: {self.colors.BG_DARKEST};
                color: {self.colors.TEXT_PRIMARY};
                border: 1px solid {self.colors.BORDER};
                border-radius: 4px;
                padding: 10px;
                font-family: 'Consolas', 'Courier New', monospace;
                font-size: 12px;
                line-height: 1.5;
            }}
        """)
        log_layout.addWidget(self.log_text)

        # Начальное сообщение
        self.log_text.append("═" * 60)
        self.log_text.append("Linx2DOI - готов к работе")
        self.log_text.append("═" * 60)
        self.log_text.append("📧 Укажите email для NCBI API")
        self.log_text.append("📄 Выберите файл .docx с библиографией")
        self.log_text.append("🚀 Нажмите 'Запустить обработку'")
        self.log_text.append("")

        self.tabs.addTab(log_tab, "Лог обработки")

        layout.addWidget(self.tabs)

    def toggle_theme(self):
        """Переключение между светлой и темной темой"""
        # Переключаем тему
        self.current_theme = "dark" if self.current_theme == "light" else "light"

        # Обновляем палитру цветов
        self.colors = LightColors if self.current_theme == "light" else AppColors

        # Сохраняем выбор в настройках
        self.settings.setValue("theme", self.current_theme)

        # Применяем новую тему
        self.apply_theme()

        # Обновляем текст кнопки
        theme_icon = "🌙" if self.current_theme == "light" else "☀️"
        theme_text = "Темная" if self.current_theme == "light" else "Светлая"
        self.theme_btn.setText(f"{theme_icon} {theme_text}")

        # Пересоздаем интерфейс для полного применения темы
        # Сохраняем текущие данные
        current_input = self.input_path
        current_output = self.output_path
        current_items = self.all_items
        current_table_items = self.current_items

        # Очищаем центральный виджет
        central_widget = QWidget()
        self.setCentralWidget(central_widget)

        # Пересоздаем интерфейс
        self.init_ui()

        # Восстанавливаем данные
        self.input_path = current_input
        self.output_path = current_output
        self.all_items = current_items
        self.current_items = current_table_items

        # Если были данные, восстанавливаем их в таблице
        if current_items:
            from ..core.html_generator import HTMLGenerator
            table_data = HTMLGenerator.generate_table_data(current_items)
            self.update_table_data(table_data)
            self.process_btn.setEnabled(True)
            self.apply_manual_doi_btn.setEnabled(True)
            self.open_btn.setEnabled(True)
            self.export_ris_btn.setEnabled(True)
            self.collect_publications_btn.setEnabled(True)
            if current_input:
                self.file_label.setText(f"📄 {current_input.name}")

    def show_help(self):
        """Показ справки"""
        help_text = """
        <h3 style='color: #3d95ed;'>Как использовать Linx2DOI:</h3>
        <ol style='line-height: 1.8;'>
        <li>Укажите ваш <b>email</b> для доступа к NCBI API</li>
        <li>Нажмите <b>"📁 Выбрать"</b> и выберите файл .docx со списком статей</li>
        <li>Нажмите <b>"🚀 Запустить обработку"</b> для автоматического поиска DOI</li>
        <li>Для статей без DOI введите DOI вручную в столбце <b>"Ручной DOI"</b></li>
        <li>Нажмите <b>"💾 Применить ручные DOI"</b> для обработки записей с ручными DOI</li>
        <li>Выберите статьи галочками и нажмите <b>"🔄 Парсить выбранные"</b> для повторной обработки</li>
        <li>Используйте вкладку <b>"Лог обработки"</b> для просмотра деталей</li>
        </ol>
        <p style='color: #969696; font-size: 13px;'><b>Формат документа:</b> Нумерованный список вида:<br>
        <code>1. Название статьи https://ссылка-на-статью</code></p>
        """
        msg = QMessageBox(self)
        msg.setWindowTitle("Справка - Linx2DOI")
        msg.setTextFormat(Qt.TextFormat.RichText)
        msg.setText(help_text)
        msg.exec()

    # === Обработчики событий ===

    def select_file(self):
        """Выбор файла .docx"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Выберите файл .docx",
            str(Path.home()),
            "Документы Word (*.docx)"
        )
        if file_path:
            self.input_path = Path(file_path)
            self.file_label.setText(f"📄 {self.input_path.name}")
            self.file_label.setStyleSheet(f"""
                QLabel {{
                    background-color: {self.colors.PRIMARY_DARK};
                    border: 1px solid {self.colors.PRIMARY};
                    border-radius: 4px;
                    padding: 6px 10px;
                    color: white;
                    font-size: 12px;
                    font-weight: 600;
                }}
            """)
            self.process_btn.setEnabled(True)
            self.apply_manual_doi_btn.setEnabled(False)
            self.open_btn.setEnabled(False)
            self.export_ris_btn.setEnabled(False)
            self.collect_publications_btn.setEnabled(False)
            self.log_text.append(f"\n✅ Выбран файл: {self.input_path.name}")
            self.results_table.setRowCount(0)
            self.all_items = []
            self.current_items = []
            self.table_items_map = {}

    @staticmethod
    def validate_email(email):
        """Валидация email"""
        pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        return re.match(pattern, email.strip()) is not None

    def start_processing(self):
        """Запуск полной обработки документа"""
        if not self.input_path:
            QMessageBox.warning(self, "Ошибка", "Сначала выберите файл .docx")
            return

        user_email = self.email_input.text().strip()
        if not user_email:
            QMessageBox.warning(self, "Ошибка", "Укажите ваш email для запросов к NCBI API")
            return

        if not self.validate_email(user_email):
            QMessageBox.warning(self, "Ошибка", "Некорректный формат email")
            return

        # Переключаемся на вкладку лога
        self.tabs.setCurrentIndex(1)

        self.settings.setValue("user_email", user_email)
        self.progress_bar.setValue(0)
        self.progress_label.setText("Обработка...")
        self.process_btn.setEnabled(False)
        self.apply_manual_doi_btn.setEnabled(False)
        self.open_btn.setEnabled(False)
        self.export_ris_btn.setEnabled(False)
        self.collect_publications_btn.setEnabled(False)
        self.log_text.clear()
        self.log_text.append("═" * 60)
        self.log_text.append("🚀 НАЧАЛО АВТОМАТИЧЕСКОЙ ОБРАБОТКИ")
        self.log_text.append("═" * 60)
        self.log_text.append(f"📧 Email: {user_email}")
        self.log_text.append(f"📄 Файл: {self.input_path.name}")
        self.log_text.append("")

        self.worker = WorkerThread(self.input_path, user_email)
        self.worker.progress.connect(self.update_progress)
        self.worker.log.connect(self.update_log)
        self.worker.finished_success.connect(self.processing_finished)
        self.worker.finished_error.connect(self.processing_error)
        self.worker.table_data.connect(self.update_table_data)
        self.worker.start()

    def process_selected(self):
        """Обработка выбранных элементов (автопоиск или ручной DOI)"""
        if not self.input_path:
            QMessageBox.warning(self, "Ошибка", "Сначала выберите файл .docx")
            return

        user_email = self.email_input.text().strip()
        if not user_email or not self.validate_email(user_email):
            QMessageBox.warning(self, "Ошибка", "Укажите корректный email")
            return

        # Собираем записи с галочками ИЛИ ручными DOI
        selected_items = []
        manual_count = 0
        auto_count = 0
        rows_to_uncheck = []  # Запоминаем строки для снятия галочек ПОСЛЕ запуска

        for row in range(self.results_table.rowCount()):
            # Получаем галочку и manual_doi
            check_widget = self.results_table.cellWidget(row, 6)
            checkbox = None
            is_checked = False

            if check_widget:
                checkbox = check_widget.layout().itemAt(0).widget()
                is_checked = checkbox and checkbox.isChecked()

            manual_doi_item = self.results_table.item(row, 4)
            manual_doi = manual_doi_item.text().strip() if manual_doi_item else ""
            has_manual_doi = manual_doi and manual_doi.startswith('10.')

            # Обрабатываем если ЛИБО галочка ЛИБО manual_doi
            if is_checked or has_manual_doi:
                item_data = self.table_items_map.get(row)
                if item_data:
                    # Определяем тип обработки
                    if has_manual_doi:
                        selected_items.append({
                            'original_item': item_data,
                            'manual_doi': manual_doi
                        })
                        manual_doi_item.setText("")  # Очищаем после сбора
                        manual_count += 1
                    else:
                        selected_items.append({
                            'original_item': item_data,
                            'manual_doi': None
                        })
                        auto_count += 1

                    # Запоминаем строку для снятия галочки (если она была)
                    if is_checked and checkbox:
                        rows_to_uncheck.append((row, checkbox))

        if not selected_items:
            QMessageBox.warning(self, "Ошибка",
                              "Не выбрано ни одной записи для обработки.\n\n"
                              "Отметьте галочками нужные записи\n"
                              "ИЛИ\n"
                              "Введите ручные DOI в колонку 'Ручной DOI'")
            return

        # Переключаемся на вкладку лога
        self.tabs.setCurrentIndex(1)

        self.progress_bar.setValue(0)
        self.progress_label.setText(f"Обработка {len(selected_items)}...")
        self.process_btn.setEnabled(False)
        self.apply_manual_doi_btn.setEnabled(False)
        self.open_btn.setEnabled(False)
        self.export_ris_btn.setEnabled(False)
        self.collect_publications_btn.setEnabled(False)
        self.log_text.append("")
        self.log_text.append("═" * 60)
        self.log_text.append(f"🔍 ОБРАБОТКА ВЫБРАННЫХ ({len(selected_items)} ЗАПИСЕЙ)")
        self.log_text.append("═" * 60)
        if manual_count > 0:
            self.log_text.append(f"💾 С ручными DOI: {manual_count}")
        if auto_count > 0:
            self.log_text.append(f"🤖 Автоматический поиск: {auto_count}")
        self.log_text.append("")

        # Проверяем наличие предыдущих результатов
        previous_items = getattr(self, 'all_items', [])

        self.worker = WorkerThread(self.input_path, user_email, selected_items,
                                   previous_items=previous_items)
        self.worker.progress.connect(self.update_progress)
        self.worker.log.connect(self.update_log)
        self.worker.finished_success.connect(self.processing_finished)
        self.worker.finished_error.connect(self.processing_error)
        self.worker.table_data.connect(self.update_table_data)
        self.worker.start()

        # Снимаем галочки после запуска worker
        for row, checkbox in rows_to_uncheck:
            checkbox.setChecked(False)

    def update_table_data(self, table_data):
        """Обновление таблицы с результатами"""
        self.results_table.setRowCount(len(table_data))
        self.table_items_map = {}

        for row, row_data in enumerate(table_data):
            self.table_items_map[row] = row_data[7]

            # № (оригинальный порядковый номер)
            item = QTableWidgetItem(row_data[0])
            item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            item.setFlags(item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            item.setForeground(QBrush(QColor(self.colors.TEXT_SECONDARY)))
            self.results_table.setItem(row, 0, item)

            # Название (гиперссылка на URL статьи)
            title_item = QTableWidgetItem(row_data[1])
            title_item.setData(Qt.ItemDataRole.UserRole + 1, True)
            title_item.setData(Qt.ItemDataRole.UserRole + 2, self.table_items_map[row]['url'])
            title_item.setForeground(QBrush(QColor(self.colors.PRIMARY_LIGHT)))
            title_item.setFont(QFont("Arial", 10, QFont.Weight.Bold))
            title_item.setFlags(title_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            self.results_table.setItem(row, 1, title_item)

            # Заголовок статьи (гиперссылка на DOI если есть)
            article_item = QTableWidgetItem(row_data[2])
            article_item.setFlags(article_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            if row_data[4] and row_data[4].strip():
                article_item.setData(Qt.ItemDataRole.UserRole + 1, True)
                article_item.setData(Qt.ItemDataRole.UserRole + 2, row_data[4])
                article_item.setForeground(QBrush(QColor(self.colors.PRIMARY_LIGHT)))
            else:
                article_item.setForeground(QBrush(QColor(self.colors.TEXT_PRIMARY)))
            self.results_table.setItem(row, 2, article_item)

            # DOI
            doi_item = QTableWidgetItem(row_data[3])
            doi_item.setFlags(doi_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            if row_data[3] == "Данные отсутствуют":
                doi_item.setForeground(QBrush(QColor(self.colors.ERROR)))
                doi_item.setFont(QFont("Arial", 10, QFont.Weight.Bold))
            else:
                doi_item.setForeground(QBrush(QColor(self.colors.SUCCESS)))
                doi_item.setFont(QFont("Consolas", 9))
            self.results_table.setItem(row, 3, doi_item)

            # Ручной DOI (редактируемый)
            manual_doi_item = QTableWidgetItem(row_data[5])
            manual_doi_item.setFlags(manual_doi_item.flags() | Qt.ItemFlag.ItemIsEditable)
            manual_doi_item.setToolTip("Введите DOI в формате 10.xxxx/xxxx")
            manual_doi_item.setFont(QFont("Consolas", 9))
            self.results_table.setItem(row, 4, manual_doi_item)

            # Статус
            status_item = QTableWidgetItem(row_data[6])
            status_item.setFlags(status_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            status_item.setFont(QFont("Arial", 10, QFont.Weight.Bold))
            if not row_data[7].get('is_cited', True):
                status_item.setForeground(QBrush(QColor("#9e9e9e")))  # Серый
            elif "НЕТ ДАННЫХ" in row_data[6]:
                status_item.setForeground(QBrush(QColor(self.colors.ERROR)))
            elif "ДУБЛЬ" in row_data[6]:
                status_item.setForeground(QBrush(QColor(self.colors.WARNING)))
            elif "РУЧНОЙ DOI" in row_data[6]:
                status_item.setForeground(QBrush(QColor(self.colors.SUCCESS)))
            else:
                status_item.setForeground(QBrush(QColor(self.colors.INFO)))
            self.results_table.setItem(row, 5, status_item)

            # Чекбокс для выбора
            check_widget = QWidget()
            check_layout = QHBoxLayout(check_widget)
            check_layout.setContentsMargins(0, 0, 0, 0)
            check_layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
            checkbox = QCheckBox()
            check_layout.addWidget(checkbox)
            self.results_table.setCellWidget(row, 6, check_widget)

        self.apply_manual_doi_btn.setEnabled(len(table_data) > 0)

        # Автоматически переключаемся на вкладку результатов после обновления
        self.tabs.setCurrentIndex(0)

    def select_all_items(self):
        """Выбрать все элементы"""
        for row in range(self.results_table.rowCount()):
            check_widget = self.results_table.cellWidget(row, 6)
            if check_widget:
                checkbox = check_widget.layout().itemAt(0).widget()
                if checkbox:
                    checkbox.setChecked(True)

    def deselect_all_items(self):
        """Снять выбор со всех элементов"""
        for row in range(self.results_table.rowCount()):
            check_widget = self.results_table.cellWidget(row, 6)
            if check_widget:
                checkbox = check_widget.layout().itemAt(0).widget()
                if checkbox:
                    checkbox.setChecked(False)

    def on_header_clicked(self, logical_index):
        """Обработка клика по заголовку колонки"""
        if logical_index == 6:  # Колонка "✓"
            # Проверяем, все ли элементы выбраны
            all_checked = True
            for row in range(self.results_table.rowCount()):
                check_widget = self.results_table.cellWidget(row, 6)
                if check_widget:
                    checkbox = check_widget.layout().itemAt(0).widget()
                    if checkbox and not checkbox.isChecked():
                        all_checked = False
                        break

            # Переключаем состояние всех галочек
            if all_checked:
                self.deselect_all_items()
            else:
                self.select_all_items()

    @pyqtSlot(int, str)
    def update_progress(self, value, status_text):
        """Обновление прогресса"""
        self.progress_bar.setValue(value)
        self.progress_label.setText(status_text)

    @pyqtSlot(str)
    def update_log(self, message):
        """Обновление лога"""
        self.log_text.append(message)
        self.log_text.verticalScrollBar().setValue(self.log_text.verticalScrollBar().maximum())

    @pyqtSlot(str, list)
    def processing_finished(self, output_path, items):
        """Обработка завершена успешно"""
        self.output_path = Path(output_path)
        self.progress_bar.setValue(100)
        self.progress_label.setText("✅ Готово!")

        self.all_items = items
        self.current_items = items

        total_items = len(items)
        with_apa = sum(1 for item in items if item.get('apa_citation'))
        no_data = sum(1 for item in items if not item.get('has_data', True))
        manual_dois = sum(1 for item in items if item.get('manual_doi'))

        self.log_text.append("")
        self.log_text.append("═" * 60)
        self.log_text.append("✅ ОБРАБОТКА ЗАВЕРШЕНА")
        self.log_text.append("═" * 60)
        self.log_text.append(f"📊 СТАТИСТИКА:")
        self.log_text.append(f"   📄 Всего: {total_items}")
        self.log_text.append(f"   📚 С цитатами: {with_apa}")
        self.log_text.append(f"   ⚠️ Без данных: {no_data}")
        self.log_text.append(f"   💾 Ручных DOI: {manual_dois}")
        self.log_text.append(f"")
        self.log_text.append(f"💾 Файл: {self.output_path.name}")

        self.process_btn.setEnabled(True)
        self.apply_manual_doi_btn.setEnabled(True)
        self.open_btn.setEnabled(True)
        self.export_ris_btn.setEnabled(True)
        self.collect_publications_btn.setEnabled(True)

        # Переключаемся на вкладку результатов
        self.tabs.setCurrentIndex(0)

        QMessageBox.information(
            self,
            "Готово",
            f"✅ Обработка завершена!\n\n"
            f"Всего: {total_items}\n"
            f"С цитатами: {with_apa}\n"
            f"Без данных: {no_data}\n"
            f"Ручных DOI: {manual_dois}\n\n"
            f"Файл: {self.output_path.name}"
        )

    @pyqtSlot(str)
    def processing_error(self, error_message):
        """Обработка ошибки"""
        self.progress_label.setText("❌ Ошибка")
        self.log_text.append(f"\n{error_message}")
        self.process_btn.setEnabled(True)
        self.apply_manual_doi_btn.setEnabled(True)
        QMessageBox.critical(self, "Ошибка",
                             error_message[:400] + "..." if len(error_message) > 400 else error_message)

    def open_result(self):
        """Открытие результата в браузере"""
        if not self.output_path or not self.output_path.exists():
            QMessageBox.warning(self, "Ошибка", "Результат ещё не сгенерирован")
            return
        try:
            if sys.platform == "win32":
                os.startfile(str(self.output_path))
            elif sys.platform == "darwin":
                os.system(f'open "{self.output_path}"')
            else:
                os.system(f'xdg-open "{self.output_path}"')
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось открыть файл:\n{str(e)}")

    def export_to_ris(self):
        """Экспорт результатов в формат RIS для Mendeley"""
        if not self.all_items or len(self.all_items) == 0:
            QMessageBox.warning(
                self,
                "Нет данных",
                "Нет обработанных данных для экспорта.\n"
                "Сначала запустите обработку документа."
            )
            return

        # Подсчет элементов с DOI
        items_with_doi = [item for item in self.all_items if item.get('dois') and len(item['dois']) > 0]

        if len(items_with_doi) == 0:
            QMessageBox.warning(
                self,
                "Нет DOI",
                "Ни один элемент не содержит DOI.\n"
                "Невозможно создать RIS файл."
            )
            return

        # Диалог сохранения файла
        default_name = "bibliography.ris"
        if self.input_path:
            default_name = self.input_path.stem + ".ris"

        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Сохранить RIS файл",
            str(Path.home() / default_name),
            "RIS файлы (*.ris);;Все файлы (*.*)"
        )

        if not file_path:
            return

        # Генерация и сохранение RIS файла
        try:
            source_filename = self.input_path.name if self.input_path else "document"
            success = RISExporter.save_ris_file(
                self.all_items,
                file_path,
                source_filename
            )

            if success:
                QMessageBox.information(
                    self,
                    "Экспорт завершен",
                    f"✅ RIS файл успешно сохранен!\n\n"
                    f"Файл: {Path(file_path).name}\n"
                    f"Экспортировано записей: {len(items_with_doi)}\n\n"
                    f"Теперь вы можете импортировать файл в:\n"
                    f"• Mendeley\n"
                    f"• Zotero\n"
                    f"• EndNote\n"
                    f"• Papers"
                )
                self.log_text.append(f"\n✅ RIS файл экспортирован: {file_path}")
                self.log_text.append(f"   Экспортировано записей: {len(items_with_doi)}")
            else:
                QMessageBox.critical(
                    self,
                    "Ошибка экспорта",
                    "Не удалось сохранить RIS файл.\n"
                    "Проверьте права доступа к папке."
                )
        except Exception as e:
            QMessageBox.critical(
                self,
                "Ошибка",
                f"Произошла ошибка при экспорте:\n{str(e)}"
            )
            self.log_text.append(f"\n❌ Ошибка экспорта RIS: {str(e)}")

    def collect_publications(self):
        """Собрать публикации в docx файл с APA форматированием"""
        if not self.all_items or len(self.all_items) == 0:
            QMessageBox.warning(
                self,
                "Нет данных",
                "Нет обработанных данных для сбора.\n"
                "Сначала запустите обработку документа."
            )
            return

        # Фильтруем записи: исключаем "ОТСУТСТВУЕТ В ТЕКСТЕ"
        filtered_items = []
        for item in self.all_items:
            is_cited = item.get('is_cited', True)
            if is_cited:  # Включаем только цитируемые записи
                filtered_items.append(item)

        if len(filtered_items) == 0:
            QMessageBox.warning(
                self,
                "Нет данных",
                "Все записи имеют статус 'ОТСУТСТВУЕТ В ТЕКСТЕ'.\n"
                "Нечего добавлять в список публикаций."
            )
            return

        # Диалог сохранения файла
        default_name = "publications.docx"
        if self.input_path:
            default_name = self.input_path.stem + "_publications.docx"

        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Сохранить список публикаций",
            str(Path.home() / default_name),
            "Word файлы (*.docx);;Все файлы (*.*)"
        )

        if not file_path:
            return

        # Создание docx файла
        try:
            doc = Document()

            # Настройка стилей
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)

            # Отслеживаем дубликаты DOI
            doi_first_occurrence = {}

            for item in filtered_items:
                # Получаем исходный номер записи
                original_index = item.get('original_index', 1)

                # Определяем статус записи
                has_data = item.get('has_data', False)
                manual_doi = item.get('manual_doi')
                apa_citation = item.get('apa_citation')

                # Проверка на дубликат
                is_duplicate = False
                duplicate_index = None

                if item['dois'] and len(item['dois']) > 0:
                    primary_doi = item['dois'][0]  # Берем первый DOI
                    if primary_doi in doi_first_occurrence:
                        # Это дубликат
                        is_duplicate = True
                        duplicate_index = doi_first_occurrence[primary_doi]
                    else:
                        # Первое вхождение
                        doi_first_occurrence[primary_doi] = original_index

                # Создаем параграф с номером (как обычный текст, не встроенная нумерация)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Добавляем номер из исходного документа
                run_number = p.add_run(f"{original_index}. ")
                run_number.font.name = 'Times New Roman'
                run_number.font.size = Pt(12)

                if is_duplicate:
                    # ДУБЛЬ - только текст "ДУБЛЬ № X"
                    run_text = p.add_run(f"ДУБЛЬ № {duplicate_index}")
                    run_text.font.name = 'Times New Roman'
                    run_text.font.size = Pt(12)

                elif not has_data:
                    # НЕТ ДАННЫХ - название с гиперссылкой
                    title = item.get('title', 'Без названия')
                    url = item.get('url', '')

                    # Добавляем название как гиперссылку
                    if url:
                        # В python-docx гиперссылки добавляются через XML
                        # Для упрощения просто добавим текст с URL
                        run_text = p.add_run(f"{title} ")
                        run_text.font.name = 'Times New Roman'
                        run_text.font.size = Pt(12)

                        # Добавляем гиперссылку
                        add_hyperlink(p, url, url)
                    else:
                        run_text = p.add_run(title)
                        run_text.font.name = 'Times New Roman'
                        run_text.font.size = Pt(12)

                else:
                    # Есть данные - полная APA цитата
                    if apa_citation:
                        # Убираем HTML теги из APA цитаты
                        clean_apa = re.sub(r'<[^>]+>', '', apa_citation)
                        run_text = p.add_run(clean_apa)
                        run_text.font.name = 'Times New Roman'
                        run_text.font.size = Pt(12)
                    else:
                        # Если нет APA, используем название
                        title = item.get('title', 'Без названия')
                        run_text = p.add_run(title)
                        run_text.font.name = 'Times New Roman'
                        run_text.font.size = Pt(12)

            # Сохраняем документ
            doc.save(file_path)

            QMessageBox.information(
                self,
                "Список публикаций создан",
                f"✅ Документ успешно сохранен!\n\n"
                f"Файл: {Path(file_path).name}\n"
                f"Включено записей: {len(filtered_items)}\n"
                f"(Исключены записи со статусом 'ОТСУТСТВУЕТ В ТЕКСТЕ')"
            )
            self.log_text.append(f"\n✅ Список публикаций создан: {file_path}")
            self.log_text.append(f"   Включено записей: {len(filtered_items)}")

        except Exception as e:
            QMessageBox.critical(
                self,
                "Ошибка",
                f"Произошла ошибка при создании файла:\n{str(e)}"
            )
            self.log_text.append(f"\n❌ Ошибка создания списка публикаций: {str(e)}")

def add_hyperlink(paragraph, url, text):
    """Добавить гиперссылку в параграф docx

    Args:
        paragraph: Параграф docx
        url: URL для гиперссылки
        text: Текст гиперссылки
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Получаем ID relationship
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    # Создаем элемент гиперссылки
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Создаем новый run
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Стиль гиперссылки (синий, подчеркнутый)
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000FF')
    rPr.append(c)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)

    # Создаем текстовый элемент
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)

    # Добавляем гиперссылку в параграф
    paragraph._p.append(hyperlink)
