"""
Парсинг документов Word для извлечения списка статей
"""
import re
from docx import Document
from docx.oxml.ns import qn


class DocumentParser:
    """Парсер для извлечения статей из .docx документов"""

    @staticmethod
    def is_bibliography_section(text):
        """Проверяет, является ли текст заголовком раздела библиографии

        Args:
            text: Текст параграфа

        Returns:
            bool: True если это заголовок раздела литературы
        """
        if not text:
            return False

        text_lower = text.strip().lower()

        bibliography_keywords = [
            'references',
            'reference list',
            'список литературы',
            'литература',
            'библиография',
            'bibliography',
            'works cited',
            'источники',
            'used literature',
            'использованная литература'
        ]

        for keyword in bibliography_keywords:
            if keyword in text_lower:
                if len(text.strip()) < 100:
                    return True

        return False

    @staticmethod
    def extract_isbn_from_text(text):
        """Извлекает ISBN из текста

        Поддерживает форматы:
        - ISBN-13: 978-0-123-45678-9
        - ISBN-10: 0-123-45678-9
        - С префиксом: ISBN: 978-0815344322

        Args:
            text: Текст для поиска ISBN

        Returns:
            str или None: Найденный ISBN (очищенный) или None
        """
        if not text:
            return None

        # Паттерн для ISBN-13 (начинается с 978 или 979)
        isbn13_pattern = r'(?:ISBN[:\s]*)?(97[89][-\s]?\d{1,5}[-\s]?\d{1,7}[-\s]?\d{1,7}[-\s]?\d)'
        # Паттерн для ISBN-10
        isbn10_pattern = r'(?:ISBN[:\s]*)?(\d[-\s]?\d{3}[-\s]?\d{5}[-\s]?\d{3}[-\s][\dX])'

        # Сначала ищем ISBN-13
        match = re.search(isbn13_pattern, text, re.IGNORECASE)
        if match:
            isbn = re.sub(r'[-\s]', '', match.group(1))
            if len(isbn) == 13:
                return isbn

        # Затем ISBN-10
        match = re.search(isbn10_pattern, text, re.IGNORECASE)
        if match:
            isbn = re.sub(r'[-\s]', '', match.group(1))
            if len(isbn) == 10:
                return isbn

        return None

    @staticmethod
    def extract_url_from_text(text):
        """Извлекает URL из текста (не из гиперссылок Word)

        Поддерживает форматы:
        - https://doi.org/...
        - doi:...
        - http://...

        Args:
            text: Текст для поиска URL

        Returns:
            str или None: Найденный URL или None
        """
        # Сначала ищем doi:номер
        doi_pattern = r'(doi:\s*10\.\d{4,9}/[-._;()/:A-Za-z0-9]+)'
        doi_match = re.search(doi_pattern, text, re.IGNORECASE)
        if doi_match:
            doi_url = doi_match.group(1).strip()
            # Нормализуем doi: в https://doi.org/
            if doi_url.lower().startswith('doi:'):
                doi_number = doi_url[4:].strip()
                return f'https://doi.org/{doi_number}'

        # Затем ищем обычные http/https URL
        url_pattern = r'(https?://[^\s<>"{}|\\^`\[\]]+)'
        matches = re.findall(url_pattern, text)
        if matches:
            url = matches[0].strip()
            url = re.sub(r'[.,;:)]+$', '', url)
            return url
        return None

    @staticmethod
    def extract_hyperlinks_from_paragraph(paragraph, doc):
        """Извлекает URL из гиперссылок в параграфе

        Args:
            paragraph: Объект параграфа python-docx
            doc: Document объект для доступа к relationships

        Returns:
            list: Список URL из гиперссылок
        """
        hyperlinks = []

        try:
            rels = doc.part.rels

            for hyperlink in paragraph._element.xpath('.//w:hyperlink'):
                rid = hyperlink.get(qn('r:id'))
                if rid and rid in rels:
                    target_url = rels[rid].target_ref
                    if target_url:
                        # ✅ Преобразуем doi: в https://doi.org/
                        if target_url.lower().startswith('doi:'):
                            doi_number = target_url[4:].strip()
                            target_url = f'https://doi.org/{doi_number}'
                        hyperlinks.append(target_url)
        except Exception:
            pass

        return hyperlinks

    @staticmethod
    def normalize_title(title):
        """Нормализует заголовок для сравнения"""
        clean = re.sub(r'[^\w\s]', ' ', title.lower())
        clean = re.sub(r'\s+', ' ', clean).strip()
        return clean[:50] if len(clean) > 50 else clean

    @staticmethod
    def clean_url(url):
        """Исправляет пробелы внутри URL"""
        url = re.sub(r'\s*([/:])\s*', r'\1', url)
        url = url.strip().rstrip('.,;:')
        return url

    @staticmethod
    def is_incomplete_url(url):
        """Проверяет, является ли URL неполным/усеченным"""
        if not url:
            return True

        if 'researchgate.net/publication/' in url.lower():
            parts = url.split('/publication/')
            if len(parts) > 1:
                publication_part = parts[1].strip()
                if not publication_part or publication_part == '/':
                    return True

        if 'pubmed.ncbi.nlm.nih.gov/' in url.lower():
            if url.rstrip('/').endswith('pubmed.ncbi.nlm.nih.gov'):
                return True

        if 'pmc.ncbi.nlm.nih.gov/articles/' in url.lower():
            if url.rstrip('/').endswith('/articles'):
                return True

        return False

    @staticmethod
    def expand_citation_range(range_str):
        """Раскрывает диапазон '5-7' в [5, 6, 7]"""
        range_str = range_str.strip()

        if '-' in range_str:
            parts = range_str.split('-')
            if len(parts) != 2:
                return []

            try:
                start = int(parts[0].strip())
                end = int(parts[1].strip())

                if start > end:
                    return []

                if end - start > 100:
                    return []

                return list(range(start, end + 1))
            except ValueError:
                return []
        else:
            try:
                return [int(range_str)]
            except ValueError:
                return []

    @staticmethod
    def collect_citations_from_text(doc):
        """Собирает все номера цитирований из текста документа"""
        cited_numbers = set()

        citation_pattern = r'\[(\d+(?:\s*[-,]\s*\d+)*(?:\s*,\s*\d+(?:\s*-\s*\d+)*)*)\]'

        for paragraph in doc.paragraphs:
            text = paragraph.text
            matches = re.findall(citation_pattern, text)

            for match in matches:
                parts = match.split(',')
                for part in parts:
                    part = part.strip()
                    numbers = DocumentParser.expand_citation_range(part)
                    cited_numbers.update(numbers)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        matches = re.findall(citation_pattern, text)

                        for match in matches:
                            parts = match.split(',')
                            for part in parts:
                                part = part.strip()
                                numbers = DocumentParser.expand_citation_range(part)
                                cited_numbers.update(numbers)

        return cited_numbers

    @staticmethod
    def extract_item_from_text(text, original_index):
        """Извлекает информацию о статье из текста

        ✅ ИЗМЕНЕНО: URL теперь опционален (книги без URL будут парситься)
        ✅ ИЗМЕНЕНО: Извлечение ISBN из текста
        """
        if not text.strip():
            return None

        # ✅ Паттерн: номер. название [URL] — URL опционален
        pattern = r'^\s*\d+\.\s*(.+?)(?:\s+(https?://.+))?$'
        match = re.match(pattern, text.strip())
        if not match:
            return None

        title_raw = match.group(1).strip()
        url_raw = match.group(2) if match.group(2) else ''

        # Если URL не найден в паттерне, ищем в тексте
        if not url_raw:
            url_raw = DocumentParser.extract_url_from_text(text)

        url = DocumentParser.clean_url(url_raw) if url_raw else ''

        # ✅ Извлечение ISBN
        isbn = DocumentParser.extract_isbn_from_text(text)

        # Удаляем URL из заголовка если он там остался
        if url and url in title_raw:
            title_raw = title_raw.replace(url, '').strip()

        # Очистка заголовка от точек в конце
        title_clean = re.sub(r'[\s\.]*\.{2,}[\s\.]*$', '', title_raw.rstrip(' \t.')).strip()
        if not title_clean:
            title_clean = title_raw.strip()

        return {
            'original_index': original_index,
            'title': title_clean,
            'url': url,
            'isbn': isbn,  # ✅ ДОБАВЛЕНО
            'normalized_title': DocumentParser.normalize_title(title_clean),
            'text': text.strip()
        }

    @staticmethod
    def get_paragraph_number(paragraph):
        """Извлекает номер из встроенной нумерации Word"""
        try:
            numPr = paragraph._element.pPr.numPr if paragraph._element.pPr is not None else None
            if numPr is not None:
                ilvl_element = numPr.ilvl
                if ilvl_element is not None:
                    return True
        except (AttributeError, KeyError):
            pass
        return None

    @staticmethod
    def collect_items_from_document(doc, analyze_citations=True):
        """Собирает все элементы из документа с сохранением порядка

        Args:
            doc: Document объект
            analyze_citations: Если True - анализировать цитирования в тексте

        Returns:
            list: Список элементов с флагом is_cited
        """
        cited_numbers = DocumentParser.collect_citations_from_text(doc) if analyze_citations else None

        items = []
        position_counter = 1

        # ✅ Поиск раздела библиографии
        bibliography_start = None
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if DocumentParser.is_bibliography_section(text):
                bibliography_start = i
                break

        # ✅ Обрабатываем только параграфы ПОСЛЕ заголовка библиографии
        paragraphs_to_process = doc.paragraphs
        if bibliography_start is not None:
            paragraphs_to_process = doc.paragraphs[bibliography_start + 1:]

        # Обработка параграфов
        i = 0
        while i < len(paragraphs_to_process):
            paragraph = paragraphs_to_process[i]
            text = paragraph.text.strip()

            if text and re.search(r'https?://$', text):
                if i + 1 < len(paragraphs_to_process):
                    next_para = paragraphs_to_process[i + 1]
                    next_text = next_para.text.strip()
                    next_has_word_num = DocumentParser.get_paragraph_number(next_para)

                    if not next_has_word_num and next_text:
                        text = text + next_text
                        i += 1

            has_text_number = text and re.match(r'^\s*\d+\.\s', text)
            has_word_numbering = DocumentParser.get_paragraph_number(paragraph)

            hyperlinks = DocumentParser.extract_hyperlinks_from_paragraph(paragraph, doc)

            if has_text_number:
                if hyperlinks and not DocumentParser.is_incomplete_url(hyperlinks[0]):
                    title_match = re.match(r'^\s*\d+\.\s*(.+?)(?:\s+https?://.*)?$', text)
                    title = title_match.group(1).strip() if title_match else text
                    numbered_text = f"{position_counter}. {title} {hyperlinks[0]}"
                    item = DocumentParser.extract_item_from_text(numbered_text, position_counter)
                else:
                    item = DocumentParser.extract_item_from_text(text, position_counter)

                if item:
                    # ✅ Если анализ цитирований отключён (cited_numbers=None), все записи citable
                    item['is_cited'] = cited_numbers is None or position_counter in cited_numbers
                    items.append(item)
                    position_counter += 1

            elif has_word_numbering and text:
                if hyperlinks and not DocumentParser.is_incomplete_url(hyperlinks[0]):
                    title = re.sub(r'\s+https?://.*$', '', text).strip()
                    numbered_text = f"{position_counter}. {title} {hyperlinks[0]}"
                    item = DocumentParser.extract_item_from_text(numbered_text, position_counter)
                else:
                    numbered_text = f"{position_counter}. {text}"
                    item = DocumentParser.extract_item_from_text(numbered_text, position_counter)

                if item:
                    # ✅ Если анализ цитирований отключён (cited_numbers=None), все записи citable
                    item['is_cited'] = cited_numbers is None or position_counter in cited_numbers
                    items.append(item)
                    position_counter += 1

            i += 1

        # Обработка таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text.strip()

                        has_text_number = text and re.match(r'^\s*\d+\.\s', text)
                        has_word_numbering = DocumentParser.get_paragraph_number(paragraph)
                        hyperlinks = DocumentParser.extract_hyperlinks_from_paragraph(paragraph, doc)

                        if has_text_number:
                            if hyperlinks and not DocumentParser.is_incomplete_url(hyperlinks[0]):
                                title_match = re.match(r'^\s*\d+\.\s*(.+?)(?:\s+https?://.*)?$', text)
                                title = title_match.group(1).strip() if title_match else text
                                numbered_text = f"{position_counter}. {title} {hyperlinks[0]}"
                                item = DocumentParser.extract_item_from_text(numbered_text, position_counter)
                            else:
                                item = DocumentParser.extract_item_from_text(text, position_counter)

                            if item:
                                # ✅ Если анализ цитирований отключён (cited_numbers=None), все записи citable
                                item['is_cited'] = cited_numbers is None or position_counter in cited_numbers
                                items.append(item)
                                position_counter += 1

                        elif has_word_numbering and text:
                            if hyperlinks and not DocumentParser.is_incomplete_url(hyperlinks[0]):
                                title = re.sub(r'\s+https?://.*$', '', text).strip()
                                numbered_text = f"{position_counter}. {title} {hyperlinks[0]}"
                                item = DocumentParser.extract_item_from_text(numbered_text, position_counter)
                            else:
                                numbered_text = f"{position_counter}. {text}"
                                item = DocumentParser.extract_item_from_text(numbered_text, position_counter)

                            if item:
                                # ✅ Если анализ цитирований отключён (cited_numbers=None), все записи citable
                                item['is_cited'] = cited_numbers is None or position_counter in cited_numbers
                                items.append(item)
                                position_counter += 1

        return items